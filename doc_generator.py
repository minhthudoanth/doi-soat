import os
import sys
from datetime import datetime
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# --- CHUYỂN SỐ THÀNH CHỮ TIẾNG VIỆT CHUẨN ---
def num_to_vietnamese_words(number):
    try:
        n = int(round(abs(float(number))))
    except:
        return "Không đồng"
        
    if n == 0:
        return "Không đồng"
        
    units = ["", "một", "hai", "ba", "bốn", "năm", "sáu", "bảy", "tám", "chín"]
    
    def read_hundreds(num, is_highest=False):
        h = num // 100
        t = (num % 100) // 10
        u = num % 10
        res = []
        
        if h > 0 or not is_highest:
            res.append(f"{units[h]} trăm")
            
        if t > 1:
            res.append(f"{units[t]} mươi")
            if u == 1:
                res.append("mốt")
            elif u == 5:
                res.append("lăm")
            elif u > 0:
                res.append(units[u])
        elif t == 1:
            res.append("mười")
            if u == 5:
                res.append("lăm")
            elif u > 0:
                res.append(units[u])
        elif t == 0 and u > 0:
            if h > 0 or not is_highest:
                res.append(f"lẻ {units[u]}")
            else:
                res.append(units[u])
                
        return " ".join(res)

    scales = ["", "nghìn", "triệu", "tỷ", "nghìn tỷ", "triệu tỷ"]
    groups = []
    temp = n
    while temp > 0:
        groups.append(temp % 1000)
        temp //= 1000
        
    res_parts = []
    for i in reversed(range(len(groups))):
        g = groups[i]
        if g > 0:
            part = read_hundreds(g, is_highest=(i == len(groups)-1))
            res_parts.append(f"{part} {scales[i]}".strip())
            
    text = " ".join(res_parts).strip()
    text = text.replace("mươi năm", "mươi lăm")
    text = text.replace("mươi một", "mươi mốt")
    text = text[0].upper() + text[1:] + " đồng"
    return text


def set_cell_margins(cell, top=80, bottom=80, left=100, right=100):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)


def set_cell_border(cell, **kwargs):
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            b = OxmlElement(f'w:{edge}')
            for key, val in edge_data.items():
                b.set(qn(f'w:{key}'), str(val))
            tcBorders.append(b)
        else:
            b = OxmlElement(f'w:{edge}')
            b.set(qn('w:val'), 'none')
            tcBorders.append(b)
    tcPr.append(tcBorders)


# --- 1. TẠO QUYẾT ĐỊNH TRUY THU (.DOCX) ---
def generate_quyet_dinh_docx(data, output_path):
    doc = docx.Document()
    
    # Chuẩn căn lề văn bản hành chính Việt Nam (Nghị định 30/2020/NĐ-CP):
    # Khổ A4: 210mm x 297mm (8.27 in x 11.69 in)
    # Lề trái: 25.4mm (1.0 in) để bấm kim/đóng tập
    # Lề phải: 20.0mm (0.79 in)
    # Lề trên: 20.0mm (0.79 in)
    # Lề dưới: 20.0mm (0.79 in)
    # Vùng in khả dụng (printable width): 6.48 in
    for sec in doc.sections:
        sec.top_margin = Inches(0.79)
        sec.bottom_margin = Inches(0.79)
        sec.left_margin = Inches(1.0)
        sec.right_margin = Inches(0.79)

    # 1. Header Table (Quốc hiệu & Tên công ty)
    table_header = doc.add_table(rows=1, cols=2)
    table_header.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_header.autofit = False
    table_header.columns[0].width = Inches(3.15)
    table_header.columns[1].width = Inches(3.33)

    c0 = table_header.cell(0, 0)
    p0 = c0.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p0.paragraph_format.line_spacing = 1.15
    p0.paragraph_format.space_after = Pt(0)
    r0 = p0.add_run("CÔNG TY CỔ PHẦN KINGFOOD MARKET\n")
    r0.bold = True
    r0.font.size = Pt(10.5)
    r0.font.name = "Times New Roman"
    r0_sub = p0.add_run("*****")
    r0_sub.font.size = Pt(10.5)
    r0_sub.font.name = "Times New Roman"

    c1 = table_header.cell(0, 1)
    p1 = c1.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.line_spacing = 1.15
    p1.paragraph_format.space_after = Pt(0)
    r1 = p1.add_run("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM\nĐộc lập – Tự do – Hạnh phúc\n")
    r1.bold = True
    r1.font.size = Pt(10.5)
    r1.font.name = "Times New Roman"
    
    # Ngày hóa đơn
    doc_date_str = data.get('doc_date', datetime.now().strftime('%d/%m/%Y'))
    try:
        dp = doc_date_str.split('/')
        date_text = f"TP.HCM, ngày {dp[0]} tháng {dp[1]} năm {dp[2]}"
    except:
        date_text = f"TP.HCM, ngày {datetime.now().day} tháng {datetime.now().month} năm {datetime.now().year}"
        
    r1_sub = p1.add_run(f"********\n{date_text}")
    r1_sub.font.italic = True
    r1_sub.font.size = Pt(10)
    r1_sub.font.name = "Times New Roman"

    for row in table_header.rows:
        for cell in row.cells:
            set_cell_border(cell)

    # 2. Tiêu đề QUYẾT ĐỊNH
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(12)
    p_title.paragraph_format.space_after = Pt(6)
    p_title.paragraph_format.line_spacing = 1.15
    r_t1 = p_title.add_run("QUYẾT ĐỊNH\n")
    r_t1.bold = True
    r_t1.font.size = Pt(13.5)
    r_t1.font.name = "Times New Roman"
    
    w_name = data.get('warehouse_name', 'KHO MEATFISH').upper()
    month = data.get('month', '07')
    year = data.get('year', '2026')
    
    r_t2 = p_title.add_run(f"Về việc truy thu giá trị claim {w_name} tháng {month}/{year}\n-------------------------")
    r_t2.bold = True
    r_t2.font.size = Pt(11.5)
    r_t2.font.name = "Times New Roman"

    # 3. Căn cứ
    p_cc = doc.add_paragraph()
    p_cc.paragraph_format.space_before = Pt(4)
    p_cc.paragraph_format.space_after = Pt(6)
    p_cc.paragraph_format.line_spacing = 1.15
    p_cc.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r_cc = p_cc.add_run(f"- Căn cứ vào kết quả giao nhận và kiểm tra thực tế tháng {month}/{year}\n- Căn cứ kết quả đối chiếu của KFM và SCF")
    r_cc.font.italic = True
    r_cc.font.size = Pt(10.5)
    r_cc.font.name = "Times New Roman"

    # 4. Điều 1
    p_d1 = doc.add_paragraph()
    p_d1.paragraph_format.space_before = Pt(4)
    p_d1.paragraph_format.space_after = Pt(6)
    p_d1.paragraph_format.line_spacing = 1.15
    p_d1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r_d1 = p_d1.add_run("Điều 1: ")
    r_d1.bold = True
    r_d1.font.size = Pt(10.5)
    r_d1.font.name = "Times New Roman"
    r_d1_txt = p_d1.add_run("Ghi nhận truy thu/bồi hoàn số lượng và giá trị theo thông tin bên dưới và bảng kê đính kèm:")
    r_d1_txt.font.size = Pt(10.5)
    r_d1_txt.font.name = "Times New Roman"

    # 5. Bảng kê chi tiết
    qty_val = abs(float(data.get('total_qty', 0)))
    amt_val = abs(float(data.get('total_amount', 0)))
    vat_type = data.get('vat_type', 'Chưa VAT')
    words = num_to_vietnamese_words(amt_val)
    inv_list = data.get('invoices', [])

    if inv_list and len(inv_list) > 1:
        # BẢNG KÊ NHIỀU HÓA ĐƠN (6 CỘT: STT | Biên bản | SL | GT | GT (Gồm VAT) | CO)
        table_data = doc.add_table(rows=len(inv_list) + 2, cols=6)
        table_data.alignment = WD_TABLE_ALIGNMENT.CENTER
        table_data.autofit = False

        headers = ["STT", "Biên bản", "SL", "GT", "GT (Gồm VAT)", "CO"]
        col_widths = [Inches(0.45), Inches(2.65), Inches(0.55), Inches(1.05), Inches(1.05), Inches(0.73)]

        for j, h in enumerate(headers):
            cell = table_data.cell(0, j)
            cell.width = col_widths[j]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.line_spacing = 1.15
            r = p.add_run(h)
            r.bold = True
            r.font.size = Pt(9.5)
            r.font.name = "Times New Roman"
            set_cell_margins(cell, 60, 60, 60, 60)
            set_cell_border(cell, top=dict(val='single', sz='6', color='000000'),
                                  bottom=dict(val='single', sz='6', color='000000'),
                                  left=dict(val='single', sz='6', color='000000'),
                                  right=dict(val='single', sz='6', color='000000'))

        tot_sl = 0
        tot_pre = 0.0
        tot_post = 0.0

        for i, it in enumerate(inv_list):
            sl = it.get('qty', 0)
            pre = it.get('pre_tax', 0.0)
            post = it.get('post_tax', 0.0)
            tot_sl += sl
            tot_pre += pre
            tot_post += post

            row_vals = [
                str(i + 1),
                it.get('content', ''),
                f"{sl:,.0f}" if sl else "",
                f"{pre:,.0f}",
                f"{post:,.0f}",
                it.get('co_number', '')
            ]
            for j, val in enumerate(row_vals):
                cell = table_data.cell(i + 1, j)
                cell.width = col_widths[j]
                p = cell.paragraphs[0]
                p.paragraph_format.line_spacing = 1.15
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if j in [2, 3, 4] else (WD_ALIGN_PARAGRAPH.CENTER if j in [0, 5] else WD_ALIGN_PARAGRAPH.LEFT)
                r = p.add_run(val)
                r.font.size = Pt(9)
                r.font.name = "Times New Roman"
                set_cell_margins(cell, 50, 50, 60, 60)
                set_cell_border(cell, top=dict(val='single', sz='4', color='CCCCCC'),
                                      bottom=dict(val='single', sz='4', color='CCCCCC'),
                                      left=dict(val='single', sz='6', color='000000'),
                                      right=dict(val='single', sz='6', color='000000'))

        # Total row
        last_row = ["", "Total", f"{tot_sl:,.0f}" if tot_sl else "", f"{tot_pre:,.0f}", f"{tot_post:,.0f}", ""]
        for j, val in enumerate(last_row):
            cell = table_data.cell(len(inv_list) + 1, j)
            cell.width = col_widths[j]
            p = cell.paragraphs[0]
            p.paragraph_format.line_spacing = 1.15
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if j in [2, 3, 4] else (WD_ALIGN_PARAGRAPH.CENTER if j in [0, 5] else WD_ALIGN_PARAGRAPH.LEFT)
            r = p.add_run(val)
            r.bold = True
            r.font.size = Pt(9.5)
            r.font.name = "Times New Roman"
            set_cell_margins(cell, 60, 60, 60, 60)
            set_cell_border(cell, top=dict(val='single', sz='6', color='000000'),
                                  bottom=dict(val='single', sz='6', color='000000'),
                                  left=dict(val='single', sz='6', color='000000'),
                                  right=dict(val='single', sz='6', color='000000'))

        amt_val = tot_pre if vat_type == 'Chưa VAT' else tot_post
        words = num_to_vietnamese_words(amt_val)

    else:
        # BẢNG TỔNG HỢP 5 CỘT
        table_data = doc.add_table(rows=3, cols=5)
        table_data.alignment = WD_TABLE_ALIGNMENT.CENTER
        table_data.autofit = False

        headers = ["Tháng", "Tên kho", "SL Chênh lệch", "Giá trị (VNĐ)", "Ghi chú"]
        col_widths = [Inches(0.7), Inches(1.8), Inches(1.1), Inches(1.44), Inches(1.44)]

        for j, h in enumerate(headers):
            cell = table_data.cell(0, j)
            cell.width = col_widths[j]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.line_spacing = 1.15
            r = p.add_run(h)
            r.bold = True
            r.font.size = Pt(10)
            r.font.name = "Times New Roman"
            set_cell_margins(cell, 70, 70, 80, 80)
            set_cell_border(cell, top=dict(val='single', sz='6', color='000000'),
                                  bottom=dict(val='single', sz='6', color='000000'),
                                  left=dict(val='single', sz='6', color='000000'),
                                  right=dict(val='single', sz='6', color='000000'))

        row1 = [f"{month}", f"{w_name}", f"({qty_val:,.0f})", f"({amt_val:,.0f})", "Claim DC 100%"]
        for j, val in enumerate(row1):
            cell = table_data.cell(1, j)
            cell.width = col_widths[j]
            p = cell.paragraphs[0]
            p.paragraph_format.line_spacing = 1.15
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if j in [2, 3] else (WD_ALIGN_PARAGRAPH.CENTER if j in [0, 4] else WD_ALIGN_PARAGRAPH.LEFT)
            r = p.add_run(val)
            r.font.size = Pt(9.5)
            r.font.name = "Times New Roman"
            set_cell_margins(cell, 60, 60, 80, 80)
            set_cell_border(cell, top=dict(val='single', sz='4', color='CCCCCC'),
                                  bottom=dict(val='single', sz='4', color='CCCCCC'),
                                  left=dict(val='single', sz='6', color='000000'),
                                  right=dict(val='single', sz='6', color='000000'))

        row2 = ["Grand Total", "", f"({qty_val:,.0f})", f"({amt_val:,.0f})", ""]
        for j, val in enumerate(row2):
            cell = table_data.cell(2, j)
            cell.width = col_widths[j]
            p = cell.paragraphs[0]
            p.paragraph_format.line_spacing = 1.15
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if j in [2, 3] else (WD_ALIGN_PARAGRAPH.CENTER if j == 0 else WD_ALIGN_PARAGRAPH.LEFT)
            r = p.add_run(val)
            r.bold = True
            r.font.size = Pt(9.5)
            r.font.name = "Times New Roman"
            set_cell_margins(cell, 60, 60, 80, 80)
            set_cell_border(cell, top=dict(val='single', sz='6', color='000000'),
                                  bottom=dict(val='single', sz='6', color='000000'),
                                  left=dict(val='single', sz='6', color='000000'),
                                  right=dict(val='single', sz='6', color='000000'))

    # 6. Chi tiết diễn giải
    p_exp = doc.add_paragraph()
    p_exp.paragraph_format.space_before = Pt(6)
    p_exp.paragraph_format.space_after = Pt(6)
    p_exp.paragraph_format.line_spacing = 1.15
    p_exp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    r_exp = p_exp.add_run(
        f"• Chi phí {w_name} T{month}/{year} ({vat_type.lower()}):\n"
        f"  - Tổng giá trị chênh lệch kho: ({amt_val:,.0f}) VNĐ ({vat_type})\n"
        f"  - Tỷ lệ quy trách nhiệm: DC (SCF) chịu 100% giá trị.\n"
        f"  - Tổng GT truy thu SCF (100%): ({amt_val:,.0f} VNĐ) ({vat_type})\n"
        f"    (Bằng chữ: {words})"
    )
    r_exp.font.size = Pt(10.5)
    r_exp.font.name = "Times New Roman"

    # 7. Điều 2 & Điều 3
    p_d2 = doc.add_paragraph()
    p_d2.paragraph_format.space_before = Pt(4)
    p_d2.paragraph_format.space_after = Pt(10)
    p_d2.paragraph_format.line_spacing = 1.15
    p_d2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    r_d2_t = p_d2.add_run("Điều 2: ")
    r_d2_t.bold = True
    r_d2_t.font.size = Pt(10.5)
    r_d2_t.font.name = "Times New Roman"
    r_d2 = p_d2.add_run(f"Khoản truy thu này sẽ được ghi nhận vào điều chỉnh năm {year}.\n")
    r_d2.font.size = Pt(10.5)
    r_d2.font.name = "Times New Roman"
    
    r_d3_t = p_d2.add_run("Điều 3: ")
    r_d3_t.bold = True
    r_d3_t.font.size = Pt(10.5)
    r_d3_t.font.name = "Times New Roman"
    r_d3 = p_d2.add_run("Quyết định có hiệu lực kể từ ngày ký. Các Phòng Ban SCF, KFM có nghĩa vụ thực hiện theo quyết định này.")
    r_d3.font.size = Pt(10.5)
    r_d3.font.name = "Times New Roman"

    # 8. Bảng Ký Tên 2 Cột Cân Đối Hoàn Toàn
    table_sign = doc.add_table(rows=1, cols=2)
    table_sign.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_sign.autofit = False
    table_sign.columns[0].width = Inches(3.24)
    table_sign.columns[1].width = Inches(3.24)

    scf_name = data.get('representative_scf', 'Nguyễn Ngọc Xuân Quang')
    kfm_name = data.get('representative_kfm', 'NGUYỄN HOÀNG LÂM')

    # Cột Trái: Đại diện SCF
    s0 = table_sign.cell(0, 0)
    ps0 = s0.paragraphs[0]
    ps0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ps0.paragraph_format.line_spacing = 1.15
    ps0.paragraph_format.space_after = Pt(0)
    
    r_scf_t = ps0.add_run("Đại diện SCF\n")
    r_scf_t.bold = True
    r_scf_t.font.size = Pt(10.5)
    r_scf_t.font.name = "Times New Roman"
    
    r_scf_sub = ps0.add_run("(Ký, họ tên)\n\n\n\n")
    r_scf_sub.font.italic = True
    r_scf_sub.bold = False
    r_scf_sub.font.size = Pt(10)
    r_scf_sub.font.name = "Times New Roman"
    
    r_scf_n = ps0.add_run(f"{scf_name}")
    r_scf_n.bold = True
    r_scf_n.font.size = Pt(10.5)
    r_scf_n.font.name = "Times New Roman"

    # Cột Phải: Đại diện KFM
    s1 = table_sign.cell(0, 1)
    ps1 = s1.paragraphs[0]
    ps1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ps1.paragraph_format.line_spacing = 1.15
    ps1.paragraph_format.space_after = Pt(0)
    
    r_kfm_t = ps1.add_run("Đại diện KFM\n")
    r_kfm_t.bold = True
    r_kfm_t.font.size = Pt(10.5)
    r_kfm_t.font.name = "Times New Roman"
    
    r_kfm_sub = ps1.add_run("(Ký, họ tên)\n\n\n\n")
    r_kfm_sub.font.italic = True
    r_kfm_sub.bold = False
    r_kfm_sub.font.size = Pt(10)
    r_kfm_sub.font.name = "Times New Roman"
    
    r_kfm_n = ps1.add_run(f"{kfm_name}")
    r_kfm_n.bold = True
    r_kfm_n.font.size = Pt(10.5)
    r_kfm_n.font.name = "Times New Roman"

    for row in table_sign.rows:
        for cell in row.cells:
            set_cell_border(cell)

    doc.save(output_path)
    return output_path


# --- 2. TẠO ĐỀ NGHỊ THANH TOÁN (.DOCX) ---
def generate_de_nghi_thanh_toan_docx(data, output_path):
    doc = docx.Document()
    
    # Chuẩn căn lề văn bản hành chính Việt Nam (Nghị định 30/2020/NĐ-CP):
    for sec in doc.sections:
        sec.top_margin = Inches(0.79)
        sec.bottom_margin = Inches(0.79)
        sec.left_margin = Inches(1.0)
        sec.right_margin = Inches(0.79)

    # 1. Header Table
    table_header = doc.add_table(rows=1, cols=2)
    table_header.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_header.autofit = False
    table_header.columns[0].width = Inches(3.15)
    table_header.columns[1].width = Inches(3.33)

    c0 = table_header.cell(0, 0)
    p0 = c0.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p0.paragraph_format.line_spacing = 1.15
    p0.paragraph_format.space_after = Pt(0)
    r0 = p0.add_run("CÔNG TY CỔ PHẦN KINGFOOD MARKET\n")
    r0.bold = True
    r0.font.size = Pt(10.5)
    r0.font.name = "Times New Roman"
    r0_sub = p0.add_run("---------o0o---------")
    r0_sub.font.size = Pt(10.5)
    r0_sub.font.name = "Times New Roman"

    c1 = table_header.cell(0, 1)
    p1 = c1.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.line_spacing = 1.15
    p1.paragraph_format.space_after = Pt(0)
    r1 = p1.add_run("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM\nĐộc Lập – Tự Do – Hạnh Phúc\n")
    r1.bold = True
    r1.font.size = Pt(10.5)
    r1.font.name = "Times New Roman"
    
    doc_date_str = data.get('doc_date', datetime.now().strftime('%d/%m/%Y'))
    try:
        dp = doc_date_str.split('/')
        date_text = f"TP.HCM, ngày {dp[0]} tháng {dp[1]} năm {dp[2]}"
    except:
        date_text = f"TP.HCM, ngày {datetime.now().day} tháng {datetime.now().month} năm {datetime.now().year}"
        
    r1_sub = p1.add_run(f"---------o0o---------\n{date_text}")
    r1_sub.font.italic = True
    r1_sub.font.size = Pt(10)
    r1_sub.font.name = "Times New Roman"

    for row in table_header.rows:
        for cell in row.cells:
            set_cell_border(cell)

    # 2. Tiêu đề ĐỀ NGHỊ THANH TOÁN
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(14)
    p_title.paragraph_format.space_after = Pt(6)
    p_title.paragraph_format.line_spacing = 1.15
    r_t1 = p_title.add_run("ĐỀ NGHỊ THANH TOÁN\n")
    r_t1.bold = True
    r_t1.font.size = Pt(13.5)
    r_t1.font.name = "Times New Roman"
    
    w_name = data.get('warehouse_name', 'KHO SEEDLOG').upper()
    month = data.get('month', '07')
    year = data.get('year', '2026')
    
    r_t2 = p_title.add_run(f"V/v Đề nghị thanh toán tiền truy thu {w_name} tháng {month}/{year}")
    r_t2.font.italic = True
    r_t2.font.size = Pt(11)
    r_t2.font.name = "Times New Roman"

    # 3. Kính gửi
    p_kg = doc.add_paragraph()
    p_kg.paragraph_format.space_before = Pt(8)
    p_kg.paragraph_format.space_after = Pt(6)
    p_kg.paragraph_format.line_spacing = 1.15
    r_kg_t = p_kg.add_run("Kính gửi: ")
    r_kg_t.bold = True
    r_kg_t.font.size = Pt(10.5)
    r_kg_t.font.name = "Times New Roman"
    r_kg = p_kg.add_run("CÔNG TY CỔ PHẦN SEEDCOM FOOD")
    r_kg.bold = True
    r_kg.font.size = Pt(10.5)
    r_kg.font.name = "Times New Roman"

    # 4. Căn cứ & Đề nghị thanh toán
    amt_val = abs(float(data.get('total_amount', 0)))
    vat_type = data.get('vat_type', 'Chưa VAT')
    words = num_to_vietnamese_words(amt_val)
    
    p_cc = doc.add_paragraph()
    p_cc.paragraph_format.space_before = Pt(4)
    p_cc.paragraph_format.space_after = Pt(6)
    p_cc.paragraph_format.line_spacing = 1.15
    p_cc.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r_cc = p_cc.add_run(
        f"- Căn cứ vào kết quả Hủy hàng tại {w_name} trong tháng {month} năm {year}\n"
        f"- Căn cứ vào kết quả đối chiếu, kiểm tra chứng từ của KFM và SCF"
    )
    r_cc.font.size = Pt(10.5)
    r_cc.font.name = "Times New Roman"
    
    p_req = doc.add_paragraph()
    p_req.paragraph_format.space_before = Pt(4)
    p_req.paragraph_format.space_after = Pt(6)
    p_req.paragraph_format.line_spacing = 1.15
    p_req.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    r_req1 = p_req.add_run("Chúng tôi kính đề nghị quý Công ty thanh toán số tiền là: ")
    r_req1.font.size = Pt(10.5)
    r_req1.font.name = "Times New Roman"
    
    r_amt = p_req.add_run(f"{amt_val:,.0f} VNĐ ({vat_type})\n")
    r_amt.bold = True
    r_amt.font.size = Pt(10.5)
    r_amt.font.name = "Times New Roman"
    
    r_words = p_req.add_run(f"(Bằng chữ: {words})")
    r_words.font.italic = True
    r_words.font.size = Pt(10.5)
    r_words.font.name = "Times New Roman"
    
    b_acc = data.get('bank_account', '04001010091039')
    b_owner = data.get('bank_owner', 'CÔNG TY CỔ PHẦN KINGFOOD MARKET')
    b_name = data.get('bank_name', 'HANG HAI (MARITIMEBANK-MSB)')
    
    p_bank = doc.add_paragraph()
    p_bank.paragraph_format.space_before = Pt(4)
    p_bank.paragraph_format.space_after = Pt(6)
    p_bank.paragraph_format.line_spacing = 1.15
    p_bank.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    r_bank = p_bank.add_run(
        f"Số tiền trên đề nghị chuyển vào tài khoản:\n"
        f" • Số tài khoản: {b_acc}\n"
        f" • Chủ tài khoản: {b_owner}\n"
        f" • Mở tại ngân hàng: {b_name}"
    )
    r_bank.font.size = Pt(10.5)
    r_bank.font.name = "Times New Roman"

    p_close = doc.add_paragraph()
    p_close.paragraph_format.space_before = Pt(4)
    p_close.paragraph_format.space_after = Pt(12)
    p_close.paragraph_format.line_spacing = 1.15
    p_close.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r_close = p_close.add_run(
        f"Kính mong Quý Công ty vui lòng thanh toán đúng thời hạn số tiền trên.\n"
        f"Trân trọng kính chào!"
    )
    r_close.font.size = Pt(10.5)
    r_close.font.name = "Times New Roman"

    # 5. Chữ ký TM. TỔNG GIÁM ĐỐC (Dùng Table căn phải hoàn hảo, không lệch lề)
    table_sign = doc.add_table(rows=1, cols=2)
    table_sign.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_sign.autofit = False
    table_sign.columns[0].width = Inches(3.24)
    table_sign.columns[1].width = Inches(3.24)

    s1 = table_sign.cell(0, 1)
    ps1 = s1.paragraphs[0]
    ps1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ps1.paragraph_format.line_spacing = 1.15
    ps1.paragraph_format.space_after = Pt(0)
    
    kfm_name = data.get('representative_kfm', 'NGUYỄN HOÀNG LÂM')
    r_sg_t = ps1.add_run("TM. TỔNG GIÁM ĐỐC\n\n\n\n\n")
    r_sg_t.bold = True
    r_sg_t.font.size = Pt(10.5)
    r_sg_t.font.name = "Times New Roman"
    
    r_sg_n = ps1.add_run(f"{kfm_name}")
    r_sg_n.bold = True
    r_sg_n.font.size = Pt(10.5)
    r_sg_n.font.name = "Times New Roman"

    for row in table_sign.rows:
        for cell in row.cells:
            set_cell_border(cell)

    doc.save(output_path)
    return output_path
